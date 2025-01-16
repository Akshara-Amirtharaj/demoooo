import streamlit as st
from docx import Document
from datetime import datetime
import os
import platform
import subprocess


# Function to replace placeholders while maintaining the existing formatting
def replace_placeholders(doc, placeholders):
    """Replace placeholders in the document while preserving formatting."""
    for para in doc.paragraphs:
        for key, value in placeholders.items():
            if key in para.text:
                for run in para.runs:
                    if key in run.text:
                        run.text = run.text.replace(key, value)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for key, value in placeholders.items():
                    if key in cell.text:
                        cell.text = cell.text.replace(key, value)


def edit_nda_template(template_path, output_path, placeholders):
    """Edit Word document to replace placeholders."""
    try:
        doc = Document(template_path)

        # Replace placeholders
        replace_placeholders(doc, placeholders)

        # Save the modified document
        doc.save(output_path)
        return output_path
    except Exception as e:
        raise Exception(f"Error editing Word template: {e}")


def convert_to_pdf(doc_path, pdf_path):
    doc_path = os.path.abspath(doc_path)
    pdf_path = os.path.abspath(pdf_path)

    if not os.path.exists(doc_path):
        raise FileNotFoundError(f"Word document not found at {doc_path}")

    if platform.system() == "Windows":
        try:
            import comtypes.client
            import pythoncom
            pythoncom.CoInitialize()
            word = comtypes.client.CreateObject("Word.Application")
            word.Visible = False
            doc = word.Documents.Open(doc_path)
            doc.SaveAs(pdf_path, FileFormat=17)
            doc.Close()
            word.Quit()
        except Exception as e:
            raise Exception(f"Error using COM on Windows: {e}")
    else:
        try:
            subprocess.run(
                ['libreoffice', '--headless', '--convert-to', 'pdf', '--outdir', os.path.dirname(pdf_path), doc_path],
                check=True
            )
        except subprocess.CalledProcessError as e:
            raise Exception(f"Error using LibreOffice: {e}")


# Initialize session state for download visibility and file paths
if "word_file_path" not in st.session_state:
    st.session_state.word_file_path = None

if "pdf_file_path" not in st.session_state:
    st.session_state.pdf_file_path = None

if "last_inputs" not in st.session_state:
    st.session_state.last_inputs = {}


def inputs_changed(current_inputs):
    """Check if the current inputs differ from the last saved inputs."""
    return st.session_state.last_inputs != current_inputs


# Streamlit app for NDA generation
st.title("NDA Document Generator")

# Input fields for NDA
base_dir = os.path.abspath(os.path.dirname(__file__))
template_path = os.path.join(base_dir, "Non Disclosure Agreement.docx")

client_name = st.text_input("Enter Client Name:", key="client_name")
company_name = st.text_input("Enter Company Name:", key="company_name")
address = st.text_area("Enter Address:", key="address")
designation = st.text_input("Enter Designation:", key="designation")
date_field = st.date_input("Enter Date:", datetime.today(), key="date_field")

placeholders = {
    "<<Client Name>>": client_name,
    "<<Company Name>>": company_name,
    "<<Address>>": address,
    "<<Designation>>": designation,
    "<<Date>>": date_field.strftime("%d-%m-%Y"),
}
# Combine inputs into a dictionary
current_inputs = {
    "client_name": client_name,
    "company_name": company_name,
    "address": address,
    "designation": designation,
    "date_field": date_field,
}

# Check if inputs have changed
if inputs_changed(current_inputs):
    st.session_state.word_file_path = None
    st.session_state.pdf_file_path = None

# Save the current inputs to session state
st.session_state.last_inputs = current_inputs

# Generate the NDA document
if st.button("Generate NDA Document"):
    formatted_date = date_field.strftime("%d %b %Y")
    file_name = f"NDA Agreement - {client_name} {formatted_date}.docx"
    pdf_file_name = f"NDA Agreement - {client_name} {formatted_date}.pdf"
    word_output_path = os.path.join(base_dir, file_name)
    pdf_output_path = os.path.join(base_dir, pdf_file_name)

    try:
        updated_path = edit_nda_template(template_path, word_output_path, placeholders)
        pdf_generated_path = convert_to_pdf(updated_path, pdf_output_path)

        # Save paths in session state
        st.session_state.word_file_path = updated_path
        st.session_state.pdf_file_path = pdf_generated_path

        st.success("NDA Document and PDF Generated Successfully!")
    except Exception as e:
        st.error(f"An error occurred: {e}")

# Display download buttons if files exist
if st.session_state.word_file_path and st.session_state.pdf_file_path:
    with open(st.session_state.word_file_path, "rb") as word_file:
        st.download_button(
            label="Download NDA Document (Word)",
            data=word_file,
            file_name=os.path.basename(st.session_state.word_file_path),
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )

    with open(st.session_state.pdf_file_path, "rb") as pdf_file:
        st.download_button(
            label="Download NDA Document (PDF)",
            data=pdf_file,
            file_name=os.path.basename(st.session_state.pdf_file_path),
            mime="application/pdf",
        )